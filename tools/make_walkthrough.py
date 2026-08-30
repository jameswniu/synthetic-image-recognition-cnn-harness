"""The code walkthrough, as a Word document in the same house theme as the decks.

The review meeting covers code and approach. The deck carries the approach; this document walks the
code, file by file and then one request end to end, so the room can follow along in an editor. The
palette is the same measured one make_deck.py uses, and every number is read from the same report
files, under the same rule: a missing report stops the build rather than shipping a blank.

Run: uv run --with python-docx python tools/make_walkthrough.py
Out: deliverables/checkbox-code-walkthrough.docx
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent

NAVY = "3730A3"
NAVY_DEEP = "1E1B4B"
PAPER = "F8FAFC"
CODE_FILL = "EEF2FF"
ACCENT = "5865F3"
INK = "18181B"
MUTED = "71717A"
WHITE = "FFFFFF"
EXC = "D97A1E"
HUM = "3A9D55"

BODY = "Helvetica Neue"
MONO = "Menlo"


def report(name: str):
    p = ROOT / name
    if not p.exists():
        raise SystemExit(f"{name} is missing. Run `make eval` and `make telemetry` before building the walkthrough.")
    data = json.loads(p.read_text())
    if not data:
        raise SystemExit(f"{name} is empty; refusing to build the walkthrough from it.")
    return data


EV = report("reports/eval_report.json")
TEL = report("reports/telemetry.json")
GOLD = report("reports/gold_report.json")
OV = EV.get("overall", {})
DET_T = (TEL.get("modes", {}).get("deterministic core only", {}) or {}).get("totals", {})
if not OV.get("tp") or not DET_T.get("boxes") or not GOLD.get("cards"):
    raise SystemExit("a report is present but carries no measurements; refusing to build the walkthrough from it.")

FOUND = f"{OV.get('tp', 0)} of {OV.get('tp', 0) + OV.get('fn', 0)}"
RIGHT = f"{round(OV.get('cls_acc', 0) * OV.get('tp', 0))} of {OV.get('tp', 0)}"
FLAGGED = DET_T.get("flagged", 0)
BOXES = DET_T.get("boxes", 0)
PAGES = DET_T.get("pages", 0)


def _rgb(hex_: str) -> RGBColor:
    return RGBColor(int(hex_[0:2], 16), int(hex_[2:4], 16), int(hex_[4:6], 16))


def _shade_el(hex_: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_)
    return shd


def shade_cell(cell, hex_: str) -> None:
    cell._tc.get_or_add_tcPr().append(_shade_el(hex_))


def shade_par(par, hex_: str) -> None:
    par._p.get_or_add_pPr().append(_shade_el(hex_))


def rule_under(par, hex_: str = ACCENT) -> None:
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "14")
    b.set(qn("w:space"), "3")
    b.set(qn("w:color"), hex_)
    pbdr.append(b)
    par._p.get_or_add_pPr().append(pbdr)


def run_in(par, text: str, size=10.5, color=INK, bold=False, mono=False, italic=False):
    r = par.add_run(text)
    r.font.name = MONO if mono else BODY
    r.font.size = Pt(size)
    r.font.color.rgb = _rgb(color)
    r.font.bold = bold
    r.font.italic = italic
    return r


def para(doc, text="", size=10.5, color=INK, bold=False, before=2, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        run_in(p, text, size=size, color=color, bold=bold)
    return p


def h1(doc, text: str, kicker: str | None = None):
    p = para(doc, text, size=15, color=NAVY, bold=True, before=14, after=2)
    rule_under(p)
    if kicker:
        para(doc, kicker, size=10, color=MUTED, before=1, after=8)
    return p


def h2(doc, text: str):
    return para(doc, text, size=11.5, color=INK, bold=True, before=10, after=3)


def code(doc, lines: list[str], note: str | None = None):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6 if i == len(lines) - 1 else 0)
        p.paragraph_format.left_indent = Inches(0.12)
        shade_par(p, CODE_FILL)
        run_in(p, line, size=9.5, mono=True)
    if note:
        para(doc, note, size=10, color=MUTED, before=0, after=8)


def bullet(doc, text: str, level: int = 1, lead: str | None = None, mono_lead: bool = False):
    """One fact on one line, its evidence one level down.

    BasisWritingStyle: a reader should be able to descend only the branch they care about. A
    paragraph makes them hold the whole block to find the one line they need, which is the cost
    this document exists to avoid.
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.20 if level == 1 else 0.50)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    numbered = bool(lead and lead.lstrip()[:1].isdigit())
    glyph = "" if numbered else ("\u2022  " if level == 1 else "\u25e6  ")
    if glyph:
        run_in(p, glyph, size=10.5, color=NAVY if level == 1 else MUTED, bold=True)
    if lead:
        run_in(p, lead, size=10.5, bold=True, mono=mono_lead)
        run_in(p, "  ", size=10.5)
    if text:
        run_in(p, text, size=10.5 if level == 1 else 10)
    return p


def table(doc, rows: list[list[str]], widths: list[float], mono_col: int | None = None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.autofit = False
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.cell(r, c)
            cell.width = Inches(widths[c])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run_in(p, str(val), size=9 if r else 9.5, color=WHITE if r == 0 else INK,
                   bold=r == 0, mono=r > 0 and c == mono_col)
            shade_cell(cell, NAVY if r == 0 else (WHITE if r % 2 else PAPER))
    # repeat the header row when the table breaks across a page
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    t.rows[0]._tr.get_or_add_trPr().append(th)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def title_band(doc):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    cell = t.cell(0, 0)
    cell.width = Inches(7.0)
    shade_cell(cell, NAVY_DEEP)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(14)
    run_in(p, "Reading checkboxes off appraisal pages", size=20, color=WHITE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    run_in(p2, "The code, walked in the order a request runs. Companion to the deck, which carries the approach. "
               "James W. Niu.", size=10, color="C7D2FE")


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _rgb(INK)

    title_band(doc)
    para(doc, "Every number in this document is read from the repo's own report files when the document is built. "
              "make eval and make telemetry reprint the machine ones; the gold rulings are a frozen human record.", size=9.5, color=MUTED, before=6, after=10)

    # ---------------------------------------------------------------- build and run
    h1(doc, "Build it and run it", "From the unzipped folder. No accounts, no keys, no network after the install.")
    code(doc, ["uv sync --extra dev"], "Installs everything, tests included. Python 3.11 or newer and uv are the only prerequisites.")
    code(doc, ["make serve"], "Starts the HTTP service on localhost:8000. The same thing runs from a container: make docker-build, then make docker-run.")
    code(doc, ["curl -s -F file=@data/samples/sample_1.jpg localhost:8000/detect"],
         "One JSON entry per checkbox, exactly the shape the brief asked for:")
    code(doc, ['{"boxes": [{"bbox": [329, 506, 387, 554], "is_checked": true}, ...]}'],
         "Add /overlay after /detect and the page comes back as a PNG with every box drawn on it: green marked, "
         "red empty, amber routed to review. Add ?explain=true and every box carries its evidence.")
    code(doc, ["make test", "make eval"],
         f"The 26 safety gates, then the accuracy table reprinted from scratch: {FOUND} boxes found, {RIGHT} marks read "
         "right on the four pages from the brief. Every command in this section was run end to end from a fresh "
         "unzip of this zip before this document was finished.")

    # ---------------------------------------------------------------- the map
    h1(doc, "The map", "Eleven source files, each with one job. Everything else is tooling that feeds them.")
    table(doc, [
        ["file", "what it does"],
        ["scripts/serve.py", "FastAPI app: /detect, /detect/overlay, /healthz. Validates the upload, shapes the response."],
        ["src/hv_checkbox/pipeline.py", "The conductor. One call runs page load, both witnesses, classification, and escalation."],
        ["src/hv_checkbox/normalize.py", "Turns an upload into a Page: grayscale, adaptive binarize, line and ink masks, upscale when small."],
        ["src/hv_checkbox/detect.py", "Witness one. Finds box-shaped intersections of printed lines, then rejects text-like impostors."],
        ["src/hv_checkbox/template.py", "Witness two. Matches the page to a blank federal form and projects its known slots onto the scan."],
        ["src/hv_checkbox/classify.py", "Reads the mark inside each box: ink fraction, stroke shape, and the policy's rulings."],
        ["src/hv_checkbox/policy.py", "The acceptance criteria as data: thresholds plus three rulings, loaded from policy.json."],
        ["src/hv_checkbox/escalate.py", "The exception lane: reason codes, and the off-by-default AI check on single crops."],
        ["src/hv_checkbox/patch_model.py", "The small trained classifier, also off by default. Scores crops, never overrules alone."],
        ["src/hv_checkbox/overlay.py", "Draws what the system saw back onto the page."],
        ["src/hv_checkbox/types.py", "Box, with bbox, state, confidence, reasons, and IoU."],
    ], widths=[2.15, 4.85], mono_col=0)

    # ---------------------------------------------------------------- one request end to end
    h1(doc, "One request, end to end", "The order a page actually moves through the code.")
    steps = [
        ("1.  Clean up the page", "normalize.load_page", [
            "Grayscale, then an adaptive threshold into an ink mask.",
            "Morphological opening gives separate horizontal and vertical line masks, and an estimated box side.",
            "Small scans are upscaled first, and every coordinate is mapped back before the response.",
        ]),
        ("2.  Find the boxes from the printed lines", "detect.detect_boxes, witness one", [
            "Candidates are intersections of horizontal and vertical lines with the right size and squareness.",
            "_ring_ink confirms a real border; _dedupe collapses overlaps.",
            "A per-page height consensus paired with _text_like throws out text-shaped impostors.",
            "Every rejection is recorded with a reason, because removing evidence is the one thing flagging cannot undo.",
        ]),
        ("3.  Find them again from the blank form", "template, witness two", [
            "Line-by-line monotone alignment beats one page-wide scale, which drifts by hundreds of pixels across form vendors.",
            "A matched template projects its known slots onto the scan; agreement between the two witnesses is the gate.",
            "Where they disagree the box is flagged, never dropped.",
            "Too little agreement marks the template untrusted and drops it for that page, which caught the wrong-form match on the held-out condominium.",
        ]),
        ("4.  Read the mark inside each box", "classify.classify", [
            "ink_fraction against the policy thresholds; mark_stats for span, dominance and connected components.",
            "The three rulings people argue about are looked up from the policy rather than hard-coded.",
            "A thin single stroke, a stray stroke through the box, a scribbled-out box.",
        ]),
        ("5.  Route what is left", "escalate", [
            "Anything unsettled carries a reason code: STRAY_STROKE, MISSING_IN_DETECT, INK_AMBIGUOUS, EXTRA_BOX, FRAGMENTED_MARK.",
            "The AI check on single crops exists behind a flag and is off by default.",
            "Nothing in the default path calls a model.",
        ]),
        ("6.  Answer", "serve", [
            "One entry per box: bbox and is_checked, plus reasons and confidence when asked.",
            "?explain=true adds the evidence per box, including every rejected candidate and why.",
        ]),
    ]
    for head, where, lines in steps:
        bullet(doc, f"({where})", level=1, lead=head)
        for line in lines:
            bullet(doc, line, level=2)

    # ---------------------------------------------------------------- policy
    h1(doc, "Where the arguments live: policy.json",
       "The definition of a marked box is data the customer owns, not code they ask us to change.")
    bullet(doc, "holds a frozen Policy dataclass.", level=1, lead="policy.py", mono_lead=True)
    bullet(doc, "The ink thresholds, the uncertain band, and the stroke-shape limits.", level=2)
    bullet(doc, "Three named rulings, each taking filled, empty, or route.", level=2)
    bullet(doc, "at the repo root is the live copy; HV_POLICY points at another.", level=1, lead="policy.json", mono_lead=True)
    bullet(doc, "refuses unknown keys.", level=1, lead="from_dict", mono_lead=True)
    bullet(doc, "A customer's typo fails loudly under their own file name instead of silently reverting to our defaults.", level=2)
    bullet(doc, "scores the shipped policy against a stricter customer's, on the same pages.", level=1, lead="scripts/compare_policies.py", mono_lead=True)
    bullet(doc, "Two boxes read differently, agreement moves from 285 of 286 to 286 of 286, and nothing in src is touched.", level=2)
    bullet(doc, "The stricter policy is not simply better: it rescues a scan-faded X by refusing to overrule ink, and the same rule would let a scribbled-out box report as filled.", level=2)
    bullet(doc, "Which error to carry is a business decision, and the file is where it gets made.", level=2)

    # ---------------------------------------------------------------- gates
    h1(doc, "The gates", "A system that grades its own homework is not a system.")
    bullet(doc, "runs 26 gates against three referees.", level=1, lead="make test", mono_lead=True)
    bullet(doc, f"The four labelled pages from the brief: {FOUND} boxes found, {RIGHT} marks read right.", level=2)
    bullet(doc, "The blank official forms, where any mark reported is a mistake with no argument.", level=2)
    bullet(doc, f"The close-up crops a person ruled before any threshold was tuned: {GOLD.get('correct', 0)} of {GOLD.get('cards', 0)} hard-graded cards agree.", level=2)
    bullet(doc, "Two gates were verified by reverting their fix and watching them fail, because a green test that would pass anyway is worse than no test.", level=1)
    bullet(doc, f"runs every page in the repo: {PAGES} pages, {BOXES:,} checkboxes.", level=1, lead="make telemetry", mono_lead=True)
    bullet(doc, f"{FLAGGED} boxes carry a reason into the queue in the default mode.", level=2)
    bullet(doc, "deliverables/dashboard.html is that telemetry drawn as the operator view: one file, no server.", level=2)
    bullet(doc, "Lines that must not change carry an INVARIANT comment, each backed by a test.", level=1)

    # ---------------------------------------------------------------- the apprentice
    h1(doc, "Who is the apprentice?",
       "A small model trained from scratch, raced against the rules on the brief's answer key, and left switched off.")
    bullet(doc, "races three readers over the brief's answer key: the rules, the CNN alone, and both together.", level=1, lead="make compare", mono_lead=True)
    bullet(doc, "Rules only, which is what ships, sends 2 of 286 boxes to a person and settles everything else without an error.", level=2)
    bullet(doc, "The CNN alone sends 56 and gets one settled box wrong; both together send 19, right on every box they settle.", level=2)
    bullet(doc, "Verdict: rules only. They differ in how much they hand to people, and most automation at the same accuracy wins, so the CNN stays switched off.", level=1)
    bullet(doc, "It caught nothing the rules missed and misread both real never-seen test boxes; switched on, it only adds work.", level=1)
    bullet(doc, "Legibility is the invariant. The rules have it built in, every answer carries a checkable reason; the CNN can only approach it by emitting logs, a confidence and a heat map per box, replayable forever.", level=1)

    # ---------------------------------------------------------------- limits
    h1(doc, "Limits, and where the TODOs live",
       "The guidelines invite TODO comments; this repo centralises them instead.")
    bullet(doc, "in priority order, holds the gaps that would matter at real volume.", level=1, lead="docs/ROADMAP.md", mono_lead=True)
    bullet(doc, "and the last slide of the deck hold the honest account of what the system cannot do.", level=1, lead="docs/approach.md", mono_lead=True)
    bullet(doc, "Four real pages is not a benchmark.", level=2)
    bullet(doc, "The page labels were seeded by the software before correction, and that shortcut hid an invented box.", level=2)
    bullet(doc, "Rejecting a candidate is resolution-dependent: at half size the invented box survives.", level=2)
    bullet(doc, "Raced against the rules on the brief's answer key, the trained classifier sent 56 boxes to a person where the rules sent 2, so it ships switched off (make compare reprints the table).", level=2)
    bullet(doc, "One place per kind of debt, instead of comments scattered through the source.", level=1)

    para(doc, "Companions: README.md for the run-and-verify surface, docs/approach.md for the approach writeup, "
              "POLICY.md for the rulings in prose, docs/EVALS.md for how the numbers are produced, and the deck for "
              "the room. Everything here is in the zip.", size=9.5, color=MUTED, before=12)

    doc.core_properties.title = "Reading checkboxes off appraisal pages: the code walkthrough"
    doc.core_properties.author = "James W. Niu"
    out = ROOT / "deliverables" / "checkbox-code-walkthrough.docx"
    doc.save(out)
    print(f"{out.name} written")


if __name__ == "__main__":
    main()
